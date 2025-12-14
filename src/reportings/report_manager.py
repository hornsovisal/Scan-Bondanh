# Reporting System with OOP Inheritance
import os
import json
from datetime import datetime
from abc import ABC, abstractmethod
import sys

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    
from host_discovery.host_discovery import HostScanner
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Base Report Manager (Abstract Base Class)
class ReportManager(ABC):
    """Base class for all report managers - demonstrates inheritance"""
    
    def __init__(self, output_dir='reports/host_scanner'):
        self.output_dir = output_dir
        self.data = []
        self._ensure_directory()
    
    def _ensure_directory(self):
        """Encapsulation: private method to create reports directory"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def add_data(self, data_item):
        """Add data to the report"""
        self.data.append(data_item)
    
    def clear_data(self):
        """Clear all collected data"""
        self.data = []
    
    @abstractmethod
    def generate_report(self, filename=None):
        """Abstract method - must be implemented by subclasses (polymorphism)"""
        pass
    
    def _get_filename(self, prefix, extension='txt'):
        """Protected method to generate timestamped filename"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return f"{prefix}_{timestamp}.{extension}"
    
    def _create_docx_header(self, doc, title):
        """Create document header with title"""
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        para = doc.add_paragraph()
        para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Spacer

# Host Discovery Report Manager (Inheritance)
class HostReportManager(ReportManager):
    """Manages host discovery reports - inherits from ReportManager"""
    
    def __init__(self, output_dir='reports/host_scanner'):
        super().__init__(output_dir)
        self.scanner = HostScanner()
        self.port_database = self._load_port_database()
    
    def _load_port_database(self):
        """Load port information from config/default_ports.json with risk levels"""
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            project_root = os.path.dirname(os.path.dirname(current_dir))
            
            if not os.path.exists(os.path.join(project_root, "config")):
                project_root = os.getcwd()
                while project_root != "/" and not os.path.exists(os.path.join(project_root, "config")):
                    project_root = os.path.dirname(project_root)
            
            config_path = os.path.join(project_root, "config/default_ports.json")
            with open(config_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Build database with risk levels
            port_db = {}
            risk_mapping = {
                'FTP': 'High',
                'Telnet': 'High',
                'SMB': 'High',
                'RDP': 'High',
                'NetBIOS': 'High',
                'SSH': 'Medium',
                'SMTP': 'Medium',
                'HTTP': 'Medium',
                'POP3': 'Medium',
                'IMAP': 'Medium',
                'HTTP-Proxy': 'Medium',
                'HTTPS': 'Low',
                'DNS': 'Low'
            }
            
            for port_info in data.get('common_ports', []):
                port = str(port_info['port'])
                service = port_info['service']
                description = port_info['description']
                risk = risk_mapping.get(service, 'Medium')
                port_db[port] = (service, risk, description)
            
            return port_db
        except Exception as e:
            # Fallback to basic database
            return {
                '21': ('FTP', 'High', 'Transmits credentials in plain text'),
                '22': ('SSH', 'Medium', 'Remote access service'),
                '80': ('HTTP', 'Medium', 'Unencrypted web traffic'),
                '443': ('HTTPS', 'Low', 'Encrypted web service')
            }
    
    def scan_and_add(self, ip_range, method='auto', max_workers=50, alive_only=False):
        """Scan IP range and add results to report data"""
        # Perform scan once and cache results
        results = self.scanner.scan(ip_range, method=method, max_workers=max_workers, alive_only=alive_only, show_progress=True)
        
        # Convert cached results to report format
        for result in results:
            ip, status, identity, ping_time, mac, ports = result
            self.add_data({
                'ip': ip,
                'status': 'Reachable' if status == 'Alive' else 'Unreachable',
                'hostname': identity,
                'ports': ports,
                'method': method.upper()
            })
    
    def load_from_cached_scan(self):
        """Load data from scanner's cached results without re-scanning"""
        results = self.scanner.get_cached_results()
        self.clear_data()  # Clear old data
        
        for result in results:
            ip, status, identity, ping_time, mac, ports = result
            self.add_data({
                'ip': ip,
                'status': 'Reachable' if status == 'Alive' else 'Unreachable',
                'hostname': identity,
                'ports': ports,
                'method': 'CACHED'
            })
    
    def _set_cell_background(self, cell, color):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def _add_risk_indicator(self, doc, level):
        """Add colored risk indicator"""
        para = doc.add_paragraph()
        run = para.add_run(f"● {level} Risk: ")
        run.font.size = Pt(10)
        
        if level == "High":
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif level == "Medium":
            run.font.color.rgb = RGBColor(255, 165, 0)
        else:
            run.font.color.rgb = RGBColor(0, 128, 0)
        
        return para
    
    def generate_report(self, filename=None, customer_name="<Company / Client Name>", 
                       network_range="<e.g. 192.168.1.0/24>", project_by="<Your Name / Team Name>"):
        """Generate professional host discovery report matching the template"""
        if not filename:
            filename = self._get_filename('host_discovery_report', 'docx')
        
        filepath = os.path.join(self.output_dir, filename)
        reachable = [d for d in self.data if d['status'] == 'Reachable']
        unreachable = [d for d in self.data if d['status'] == 'Unreachable']
        
        # Create Word document
        doc = Document()
        
        # ===== TITLE =====
        title = doc.add_heading('NETWORK SCANNING - Host Discovery REPORT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title.runs[0].font.size = Pt(18)
        title.runs[0].font.bold = True
        title.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        title.runs[0].font.name = 'Arial'
        
        # ===== METADATA =====
        metadata = [
            f"Customer Name: {customer_name}",
            f"Network Range Scanned: {network_range}",
            f"Scan Date: {datetime.now().strftime('%d-%b-%Y')}",
            f"Report Generated: {datetime.now().strftime('%d-%b-%Y %H:%M:%S')}",
            f"Project By: {project_by}"
        ]
        
        for text in metadata:
            para = doc.add_paragraph()
            if ':' in text:
                label, value = text.split(':', 1)
                run_label = para.add_run(label + ':')
                run_label.font.bold = True
                run_label.font.size = Pt(11)
                run_label.font.name = 'Arial'
                run_value = para.add_run(value)
                run_value.font.size = Pt(11)
                run_value.font.name = 'Arial'
            else:
                run = para.add_run(text)
                run.font.size = Pt(11)
                run.font.name = 'Arial'
        
        doc.add_paragraph()
        
        # ===== 1. EXECUTIVE SUMMARY =====
        heading1 = doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
        heading1.runs[0].font.size = Pt(14)
        heading1.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        heading1.runs[0].font.name = 'Arial'
        
        para = doc.add_paragraph(
            f"An automated network host discovery and port scanning assessment was conducted using a "
            f"Python-based security tool. The scan allowed rapid identification of active hosts, "
            f"open ports, and exposed services across the target network. The goal of this assessment "
            f"was to evaluate security risks and provide recommendations for strengthening network defenses."
        )
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = 'Arial'
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
        
        # ===== 2. SCAN OBJECTIVE =====
        heading2 = doc.add_heading('2. SCAN OBJECTIVE', level=1)
        heading2.runs[0].font.size = Pt(14)
        heading2.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        heading2.runs[0].font.name = 'Arial'
        
        objectives = [
            "Automatically discover live hosts in the target network",
            "Identify open ports and running services",
            "Map network topology and device connectivity",
            "Detect potential security exposures",
            "Generate a structured security report"
        ]
        for obj in objectives:
            p = doc.add_paragraph(f"● {obj}", style='List Bullet')
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.name = 'Arial'
        doc.add_paragraph()
        
        # ===== 3. SCAN METHODOLOGY =====
        heading3 = doc.add_heading('3. SCAN METHODOLOGY (AUTOMATED)', level=1)
        heading3.runs[0].font.size = Pt(14)
        heading3.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        heading3.runs[0].font.name = 'Arial'
        
        methodology = [
            "The scan was executed using a Python-based automation script",
            "ICMP and TCP techniques were used for host discovery",
            "Multi-threaded scanning enabled fast network enumeration",
            "Service detection was carried out using known port-to-service mapping",
            "Risk levels were automatically assigned based on open ports and service data",
            "The final report was generated from the automated scan output"
        ]
        for method in methodology:
            p = doc.add_paragraph(f"● {method}", style='List Bullet')
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.name = 'Arial'
        doc.add_paragraph()
        
        # ===== 4. DETAILED PORT & HOST REPORT =====
        heading4 = doc.add_heading('4. DETAILED PORT & HOST REPORT', level=1)
        heading4.runs[0].font.size = Pt(14)
        heading4.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        heading4.runs[0].font.name = 'Arial'
        
        # Summary table
        table = doc.add_table(rows=len(reachable) + 1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header
        headers = ['Host Name', 'IP Address', 'Open Ports', 'Services Running', 'Risk Level']
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            self._set_cell_background(cell, '4472C4')  # Professional blue
            # Style header text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                    run.font.name = 'Arial'
        
        # Data rows
        for idx, host in enumerate(reachable, 1):
            cells = table.rows[idx].cells
            
            # Determine host name
            host_name = f"PC-{idx:02d}" if host['hostname'] == host['ip'] else host['hostname']
            cells[0].text = host_name
            cells[1].text = host['ip']
            
            # Parse ports
            ports_list = host['ports'].split(',') if host['ports'] != 'None' else []
            cells[2].text = ', '.join(ports_list) if ports_list else 'None'
            
            # Map ports to services using loaded database
            services = []
            risk_level = "Low"
            for port in ports_list:
                port_num = port.strip()
                if port_num in self.port_database:
                    service, level, _ = self.port_database[port_num]
                    services.append(service)
                    # Determine highest risk level
                    if level == "High":
                        risk_level = "High"
                    elif level == "Medium" and risk_level != "High":
                        risk_level = "Medium"
                else:
                    services.append(f"Port-{port_num}")
            
            cells[3].text = ', '.join(services) if services else 'N/A'
            cells[4].text = risk_level
            
            # Style all cells
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
            
            # Color code risk level
            if risk_level == "High":
                self._set_cell_background(cells[4], 'FF6B6B')
                cells[4].paragraphs[0].runs[0].font.bold = True
                cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(139, 0, 0)
            elif risk_level == "Medium":
                self._set_cell_background(cells[4], 'FFD93D')
                cells[4].paragraphs[0].runs[0].font.bold = True
                cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(184, 134, 11)
            else:
                self._set_cell_background(cells[4], '6BCF7F')
                cells[4].paragraphs[0].runs[0].font.bold = True
                cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 100, 0)
        
        doc.add_paragraph()
        
        # Risk legend
        self._add_risk_indicator(doc, "High")
        run = doc.paragraphs[-1].add_run("Public services like FTP and SMB")
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        
        self._add_risk_indicator(doc, "Medium")
        run = doc.paragraphs[-1].add_run("Web & Remote Access")
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        
        self._add_risk_indicator(doc, "Low")
        run = doc.paragraphs[-1].add_run("Normal internal services")
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        
        doc.add_paragraph()
        
        # ===== 5. PORT RISK CLASSIFICATION =====
        heading5 = doc.add_heading('5. Port Risk Classification', level=1)
        heading5.runs[0].font.size = Pt(14)
        heading5.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        heading5.runs[0].font.name = 'Arial'
        
        # Collect all unique open ports from scan results
        all_open_ports = set()
        for host in reachable:
            if host['ports'] != 'None':
                ports_list = host['ports'].split(',')
                for port in ports_list:
                    all_open_ports.add(port.strip())
        
        # Filter to only show ports that were found
        risk_data = []
        for port in sorted(all_open_ports, key=lambda x: int(x) if x.isdigit() else 0):
            if port in self.port_database:
                service, level, reason = self.port_database[port]
                risk_data.append((port, service, level, reason))
            else:
                risk_data.append((port, 'Unknown', 'Medium', 'Unidentified service'))
        
        # Create table with dynamic rows
        if risk_data:
            risk_table = doc.add_table(rows=len(risk_data) + 1, cols=3)
            risk_table.style = 'Light Grid Accent 1'
            
            # Headers
            header_labels = ['Port', 'Service', 'Risk Level\nReason']
            for idx, label in enumerate(header_labels):
                cell = risk_table.rows[0].cells[idx]
                cell.text = label
                self._set_cell_background(cell, '4472C4')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.name = 'Arial'
            
            # Data rows
            for idx, (port, service, level, reason) in enumerate(risk_data, 1):
                cells = risk_table.rows[idx].cells
                cells[0].text = port
                cells[1].text = service
                cells[2].text = f"{level}\n{reason}"
                
                # Style cells
                for cell in cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Arial'
                
                if level == "High":
                    self._set_cell_background(cells[2], 'FF6B6B')
                    cells[2].paragraphs[0].runs[0].font.bold = True
                    cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(139, 0, 0)
                elif level == "Medium":
                    self._set_cell_background(cells[2], 'FFD93D')
                    cells[2].paragraphs[0].runs[0].font.bold = True
                    cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(184, 134, 11)
                else:
                    self._set_cell_background(cells[2], '6BCF7F')
                    cells[2].paragraphs[0].runs[0].font.bold = True
                    cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 100, 0)
        else:
            doc.add_paragraph("No open ports detected during scan.")
        
        # Save document
        doc.save(filepath)
        print(f"✓ Professional DOCX Report saved to: {filepath}")
        return filepath


# Port Scan Report Manager (Inheritance)
class PortReportManager(ReportManager):
    """Manages port scanning reports - inherits from ReportManager"""
    
    def __init__(self, output_dir='reports/port_scanner'):
        super().__init__(output_dir)
        self.port_database = self._load_port_database()
        self.target_ip = None
        self.target_hostname = None
    
    def _load_port_database(self):
        """Load port information from config/default_ports.json"""
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            project_root = os.path.dirname(os.path.dirname(current_dir))
            
            if not os.path.exists(os.path.join(project_root, "config")):
                project_root = os.getcwd()
                while project_root != "/" and not os.path.exists(os.path.join(project_root, "config")):
                    project_root = os.path.dirname(project_root)
            
            config_path = os.path.join(project_root, "config/default_ports.json")
            with open(config_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Build database
            port_db = {}
            for port_info in data.get('common_ports', []):
                port_db[port_info['port']] = {
                    'service': port_info['service'],
                    'protocol': port_info['protocol'],
                    'description': port_info['description']
                }
            
            return port_db
        except Exception as e:
            print(f"Warning: Could not load port database: {e}")
            return {}
    
    def get_port_info(self, port: int):
        """Get port information from database"""
        return self.port_database.get(port, {
            'service': 'Unknown',
            'protocol': 'tcp',
            'description': 'Unknown service'
        })
    
    def classify_risk(self, port: int, banner: str = ""):
        """
        Classify port risk level based on port number and banner
        Returns (risk_level, notes)
        """
        # High-risk ports
        high_risk = {
            21: "Unencrypted FTP - credentials transmitted in clear text",
            23: "Telnet - unencrypted remote access, highly vulnerable",
            3389: "RDP exposed - common ransomware target",
            445: "SMB exposed - vulnerable to ransomware attacks",
            139: "NetBIOS exposed - information disclosure risk",
            1433: "MS SQL Server exposed - database access risk",
            3306: "MySQL exposed - database access risk",
            5432: "PostgreSQL exposed - database access risk",
            27017: "MongoDB exposed - NoSQL database access risk"
        }
        
        # Medium-risk ports
        medium_risk = {
            80: "HTTP - unencrypted web traffic, data interception possible",
            8080: "HTTP Proxy/Alt - unencrypted, often dev/staging server",
            25: "SMTP - email server, can be abused for spam/relay",
            110: "POP3 - unencrypted email retrieval",
            143: "IMAP - unencrypted email access",
            53: "DNS exposed - potential for DNS amplification attacks"
        }
        
        # Low-risk ports
        low_risk = {
            22: "SSH - encrypted, but ensure strong authentication",
            443: "HTTPS - encrypted web traffic, standard and secure",
            993: "IMAPS - encrypted email (secure)",
            995: "POP3S - encrypted email (secure)",
            587: "SMTP TLS - encrypted email submission"
        }
        
        # Determine risk level
        if port in high_risk:
            risk = "High"
            notes = high_risk[port]
        elif port in medium_risk:
            risk = "Medium"
            notes = medium_risk[port]
        elif port in low_risk:
            risk = "Low"
            notes = low_risk[port]
        else:
            risk = "Medium"
            notes = "Service identified but risk profile unknown"
        
        # Add banner-specific notes
        if banner and banner != "N/A":
            banner_lower = banner.lower()
            if "werkzeug" in banner_lower or "flask" in banner_lower:
                risk = "High"
                notes = "Development/staging server detected - should not be exposed in production"
            elif "apache" in banner_lower or "nginx" in banner_lower:
                notes += ". Web server banner detected"
            elif "ssh" in banner_lower and "openssh" in banner_lower:
                notes = "Standard SSH service - ensure key-based authentication"
        
        return risk, notes
    
    def _set_cell_background(self, cell, color):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def _sanitize_text(self, text: str) -> str:
        """Sanitize text to be XML compatible - remove control characters and NULL bytes"""
        if not text:
            return ""
        # Remove control characters and NULL bytes, keep printable characters
        import re
        # Keep only printable ASCII and common unicode characters
        sanitized = ''.join(char for char in text if ord(char) >= 32 or char in '\t\n\r')
        # Remove any NULL bytes
        sanitized = sanitized.replace('\x00', '')
        return sanitized
    
    def set_target_info(self, ip: str, hostname: str = None):
        """Set target IP and hostname for the report"""
        self.target_ip = self._sanitize_text(ip) if ip else None
        self.target_hostname = self._sanitize_text(hostname) if hostname else None
    
    def add_port_result(self, port: int, state: str, banner: str = "N/A"):
        """Add a port scan result"""
        port_info = self.get_port_info(port)
        # Sanitize banner before processing
        banner = self._sanitize_text(banner) if banner else "N/A"
        risk_level, notes = self.classify_risk(port, banner)
        
        self.add_data({
            'port': port,
            'protocol': self._sanitize_text(port_info['protocol']),
            'state': self._sanitize_text(state),
            'banner': banner if banner != "N/A" else f"N/A ({port_info['service']} Detected)",
            'risk_level': self._sanitize_text(risk_level),
            'notes': self._sanitize_text(notes)
        })
    
    def generate_report(self, filename=None, customer_name="<Company / Client Name>", 
                       project_by="Scan Bondanh Team - Horn Sovisal, Kuyseng Marakat, Chhit Sovathana"):
        """Generate professional port scan report in DOCX format"""
        if not filename:
            timestamp = datetime.now().strftime("%b-%d-%Y_%I-%M-%S_%p")
            filename = f"port_scanning_{timestamp}.docx"
        
        filepath = os.path.join(self.output_dir, filename)
        
        # Create Word document
        doc = Document()
        
        # ===== TITLE =====
        title = doc.add_heading('NETWORK SCANNING - Port Scanning REPORT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title.runs[0].font.size = Pt(18)
        title.runs[0].font.bold = True
        title.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        title.runs[0].font.name = 'Arial'
        
        # ===== METADATA =====
        metadata = [
            f"Customer Name: {customer_name}",
            f"Target Scanned: {self.target_ip}" + (f" ({self.target_hostname})" if self.target_hostname else ""),
            f"Scan Date: {datetime.now().strftime('%d-%b-%Y')}",
            f"Report Generated: {datetime.now().strftime('%d-%b-%Y %H:%M:%S')}",
            f"Project By: {project_by}"
        ]
        
        for item in metadata:
            p = doc.add_paragraph()
            key, value = item.split(':', 1)
            run1 = p.add_run(key + ':')
            run1.font.bold = True
            run1.font.name = 'Arial'
            run1.font.size = Pt(11)
            run2 = p.add_run(value)
            run2.font.name = 'Arial'
            run2.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(2)
        
        doc.add_paragraph()
        
        # ===== 1. EXECUTIVE SUMMARY =====
        heading = doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
        heading.runs[0].font.name = 'Arial'
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        open_ports = [d for d in self.data if d['state'] == 'OPEN']
        high_risk = len([d for d in open_ports if d['risk_level'] == 'High'])
        medium_risk = len([d for d in open_ports if d['risk_level'] == 'Medium'])
        low_risk = len([d for d in open_ports if d['risk_level'] == 'Low'])
        
        summary_text = (
            f"An automated port scanning assessment was conducted on the target system "
            f"{self.target_ip}. The scan identified {len(open_ports)} open port(s) out of "
            f"{len(self.data)} ports scanned. Risk analysis revealed {high_risk} high-risk, "
            f"{medium_risk} medium-risk, and {low_risk} low-risk services exposed. "
            f"The goal of this assessment was to identify potentially vulnerable services "
            f"and provide recommendations for strengthening network security."
        )
        
        p = doc.add_paragraph(summary_text)
        p.runs[0].font.name = 'Arial'
        p.runs[0].font.size = Pt(11)
        p.paragraph_format.space_after = Pt(12)
        
        doc.add_paragraph()
        
        # ===== 2. SCAN OBJECTIVE =====
        heading = doc.add_heading('2. SCAN OBJECTIVE', level=1)
        heading.runs[0].font.name = 'Arial'
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        objective_text = (
            f"To identify all open ports and running services on target {self.target_ip}, "
            f"assess their security risk levels, and provide actionable recommendations "
            f"for improving the security posture of exposed network services."
        )
        
        p = doc.add_paragraph(objective_text)
        p.runs[0].font.name = 'Arial'
        p.runs[0].font.size = Pt(11)
        p.paragraph_format.space_after = Pt(12)
        
        doc.add_paragraph()
        
        # ===== 3. DETAILED PORT SCAN RESULTS =====
        heading = doc.add_heading(f'3. Host Details: [{self.target_ip}]', level=1)
        heading.runs[0].font.name = 'Arial'
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # ===== CREATE TABLE =====
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Header row with professional styling
        hdr = table.rows[0].cells
        headers = ["Port", "Protocol", "State", "Identified Banner / Fingerprint", "Risk Level", "Notes"]
        for i, header_text in enumerate(headers):
            hdr[i].text = header_text
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Set cell background to professional blue
            self._set_cell_background(hdr[i], "4472C4")
        
        # ===== ADD DATA ROWS =====
        open_ports = [d for d in self.data if d['state'] == 'OPEN']
        
        for data in open_ports:
            row = table.add_row().cells
            row[0].text = str(data['port'])
            row[1].text = data['protocol']
            row[2].text = data['state']
            row[3].text = data['banner']
            row[4].text = data['risk_level']
            row[5].text = data['notes']
            
            # Apply font styling
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
            
            # Color-code risk level
            risk_cell = row[4]
            for paragraph in risk_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    risk_level = data['risk_level']
                    if risk_level == "High":
                        run.font.color.rgb = RGBColor(255, 107, 107)  # Red
                    elif risk_level == "Medium":
                        run.font.color.rgb = RGBColor(255, 217, 61)  # Yellow
                    elif risk_level == "Low":
                        run.font.color.rgb = RGBColor(107, 207, 127)  # Green
        
        # ===== SUMMARY =====
        doc.add_paragraph()
        summary = doc.add_paragraph()
        summary.add_run("Summary: ").bold = True
        summary.add_run(f"{len(open_ports)} open port(s) detected out of {len(self.data)} scanned.")
        for run in summary.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
        
        # Save document
        doc.save(filepath)
        
        print(f"\n✅ Professional port scan report saved:")
        print(f"   Location: {filepath}")
        print(f"   {len(open_ports)} open port(s) detected")
        
        return filepath


# Usage Example
if __name__ == "__main__":
    # Example 1: Using HostScanner directly, then generate professional report from cache
    # scanner = HostScanner()
    # sample_range = scanner.icmp.generate_ip_range("192.168.100.1", "192.168.100.214")
    
    # # Scan and display
    # scanner.display(sample_range, method='auto')
    
    # # Generate professional report from cached results (NO RE-SCANNING)
    # host_report = HostReportManager()
    # host_report.scanner = scanner  # Use the same scanner with cached results
    # host_report.load_from_cached_scan()  # Load from cache
    
    # # Generate report with custom metadata
    # host_report.generate_report(
    #     customer_name="CADT Network Security Team",
    #     network_range="172.23.3.0/24",
    #     project_by="Horn Sovisal, Kuyseng Marakat, Chhit Sovathana"
    # )
    
    # print("\n✓ Professional report generated matching security assessment template!")
    scanner


