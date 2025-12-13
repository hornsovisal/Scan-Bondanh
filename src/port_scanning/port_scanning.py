# --- Library Imports ---
from __future__ import annotations   # This will tell the hints to avoid error
import socket     # This one is for network communications
import re     # Check whether the IP is private or not 
import json   # use in scanner_result saving in JSON format(easy to read)
import csv    # use in scanner_result saving in CSV format
import sys    # use to exit the program safely(control how your program start and stop)
from pathlib import Path # use to handle file and directory paths in a clean, especially in cross-platform way
from datetime import datetime # use to get the current date
from concurrent.futures import ThreadPoolExecutor, as_completed #use to run port scans in parallel using multiple threads
from typing import List, Tuple   # this will return value is a list of (port, status)
from docx import Document # This use to save the result in docx format


# --- Color the output so that it's more readable ---
try:
    from colorama import init, Fore, Style
    init(autoreset=True)
except Exception:
    class _C:
        RESET = ""
        BRIGHT = ""
    class Fore:
        RED = GREEN = YELLOW = CYAN = RESET = ""
    Style = _C()



# --- These are common service ports used by many network services---
DEFAULT_PORTS = [
    20, 21, 22, 23, 25, 53, 67, 68, 69, 80, 88, 110, 123,
    135, 139, 143, 161, 162, 389, 443, 445, 465, 587,
    631, 636, 993, 995, 1433, 1521, 1723, 3306, 3389,
    5432, 27017
]


# --- Show the current time in the scanner_result ---
def readable_time() -> str:
    return datetime.now().strftime("%b-%d-%Y_%I-%M-%S_%p")


# =====================================================
# BASE CLASS (PARENT CLASS)
# =====================================================
class BaseScanner:
    def __init__(self, workers: int, timeout: float):
        # Private attributes
        self.__workers = workers
        self.__timeout = timeout

        # Create a main output folder to store scan results
        base_dir = Path(__file__).resolve().parents[1]  # src/
        self.__output_dir = base_dir / "reportings" / "reports"
        if not self.__output_dir.exists():
            raise FileNotFoundError(
                f"Required folder does not exist: {self.__output_dir}"
            )


# --- Methods ---
    def resolve_target(self, target: str) -> str | None:
        """
        It will convert a domain name (example: google.com) into an IP address.

        - If the domain is valid, return its IP address
        - If the domain cannot be resolved, return None

        This prevents the program from crashing when the user
        enters an invalid domain or IP.
        """
        try:
            return socket.gethostbyname(target)
        except Exception:
            return None

    def is_private_ip(self, ip: str) -> bool:
        """
        Check whether the given IP address is private or not.

        Private IP addresses are blocked to prevent scanning
        local or internal networks.
        """
        patterns = [
            r"^10\.",
            r"^192\.168\.",
            r"^127\.",
            r"^169\.254\.",
            r"^172\.(1[6-9]|2[0-9]|3[0-1])\."
        ]
        return any(re.match(p, ip) for p in patterns)


    def save_results(self, results, ip):
        """
        Save port scan results in Nmap-style summary format
        directly into:
        src/reportings/reports/
        """

    # --- Timestamp ---
        timestamp = datetime.now().strftime("%b-%d-%Y_%I-%M-%S_%p")

    # --- Reports directory (already exists) ---
        reports_dir = Path(__file__).resolve().parents[1] / "reportings" / "reports"

        report_path = reports_dir / f"port_scanning_{timestamp}.docx"

    # --- Service & Risk mapping ---
        SERVICE_MAP = {
           21: ("FTP", "High"),
           22: ("SSH", "High"),
           80: ("HTTP", "Medium"),
           443: ("HTTPS", "Low"),
           27017: ("MongoDB", "High"),
        }


    # --- Create DOCX ---
        doc = Document()
        doc.add_heading("Scan Summary (Nmap-style)", level=1)

        doc.add_paragraph("-" * 75)

    # --- Table header ---
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Host Name"
        hdr[1].text = "IP Address"
        hdr[2].text = "Open Port"
        hdr[3].text = "Service"
        hdr[4].text = "Risk"

        try:
           hostname = socket.gethostbyaddr(ip)[0]
        except Exception:
           hostname = "Unknown"

    # --- Only OPEN ports ---
        for port, status in results:
            if status != "OPEN":
               continue

        service, risk = SERVICE_MAP.get(port, ("Unknown", "Unknown"))

        row = table.add_row().cells
        row[0].text = hostname
        row[1].text = ip
        row[2].text = str(port)
        row[3].text = service
        row[4].text = risk

        doc.add_paragraph("-" * 75)

        doc.save(report_path)

        print("\n✅ Nmap-style report saved successfully:")
        print(f"   {report_path}")

    # using gettter to access to private attributes
    def get_workers(self):
        return self.__workers

    def get_timeout(self):
        return self.__timeout
    
    def print_nmap_summary(host, ip, results):
        print("\nScan Summary (Nmap-style)\n")
        print("-" * 75)

        print(
            f"{'Host Name':<15}"
            f"{'IP Address':<18}"
            f"{'Open Ports':<15}"
            f"{'Services':<18}"
            f"{'Risk'}"
        )

        print("-" * 75)

        found = False

        for port, status in results:
            if status == "OPEN":
               found = True

            # simple service mapping
            service_map = {
                80: ("HTTP", "Medium"),
                443: ("HTTPS", "Low"),
            }

            service, risk = service_map.get(port, ("Unknown", "Low"))

            print(
                f"{host:<15}"
                f"{ip:<18}"
                f"{port:<15}"
                f"{service:<18}"
                f"{risk}"
            )

        if not found:
           print("No open ports found.")

        print("-" * 75)



# =====================================================
# CHILD CLASS (INHERITANCE)
# =====================================================
class PortScanner(BaseScanner):
    def __init__(self):
        # using super(). to call the parent class constructor 
        super().__init__(workers=200, timeout=0.6)

        # using private attribute to store the list of ports to scan
        self.__ports = DEFAULT_PORTS

# --- Methods ---
    def scan_single_port(self, ip: str, port: int) -> Tuple[int, str]:
        """
        Scan a single port on the target IP

        Returns:
        - (port, "OPEN") if connection succeeds
        - (port, "CLOSED") if refused
        - (port, "FILTERED") if timed out or protected by firewall
        """
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(self.get_timeout())
        try:
            sock.connect((ip, port))
            return port, "OPEN"
        except socket.timeout:
            return port, "FILTERED"
        except ConnectionRefusedError:
            return port, "CLOSED"
        except Exception:
            return port, "CLOSED"
        finally:
            sock.close()

    def scan_ports(self, ip: str) -> List[Tuple[int, str]]:
        """
        Scan all ports using multithreading for faster performance
        """
        print(f"\nScanning {ip} ({len(self.__ports)} ports)...\n")
        results = []

        with ThreadPoolExecutor(max_workers=self.get_workers()) as executor:
            futures = [
                executor.submit(self.scan_single_port, ip, p)
                for p in self.__ports
            ]
            for future in as_completed(futures):
                results.append(future.result())

        return sorted(results, key=lambda x: x[0])

    # This is the main program flow 
    def run(self):
        """
        Main execution flow of the port scanner.
        """
        target = input("Enter target (public IP or domain): ").strip()
        ip = self.resolve_target(target)

        # Stop if the domain or IP cannot be resolved
        if not ip:
            print("❌ Cannot resolve target.")
            return

        # Block private IP addresses
        if self.is_private_ip(ip):
            print("❌ Private IP detected. Scan blocked.")
            return

        print(f"Resolved target → {ip}")

        results = self.scan_ports(ip)

        # Display results in the terminal
        print("\nPORT     STATE")
        print("--------------")

        open_found = False

        for port, status in results:
            if status == "OPEN":
               open_found = True
               print(f"{port}/tcp  open")

        if not open_found:
            print("No open ports found.")

        # Ask user whether to save results
        if input("\nSave results? (y/N): ").lower() == "y":
            self.save_results(results, ip)

        print("\nScan completed.")


# --- entry point program ---
if __name__ == "__main__":
    try:
        PortScanner().run()
    except KeyboardInterrupt:
        print("\nScan interrupted.")
        sys.exit(0)
